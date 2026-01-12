
import os
import sys
from dotenv import load_dotenv

sys.path.append(os.getcwd())
try:
    from summarizer import summarize_document, save_summary_to_pdf
except ImportError as e:
    print(f"Import Error: {e}")
    sys.exit(1)

load_dotenv()

def verify_summarizer():
    # Test file: using the EPUB from previous steps or any small file
    # If no file exists, create a dummy text file
    
    # Let's try to reuse the EPUB if it exists, otherwise create a dummy file
    epub_path = r"C:\Users\admin\Downloads\Telegram Desktop\Rocio Salas-Whalen - Weightless.epub"
    
    if os.path.exists(epub_path):
        filename = epub_path
        mime_type = "application/epub+zip"
        print(f"Testing with Real EPUB: {filename}")
    else:
        # Create dummy text file
        filename = "test_summary.txt"
        mime_type = "text/plain" # Note: our loader doesn't support text/plain explicitly yet, let's use docx mock
        # Actually our loader only supports docx/epub/pdf. Let's fail if no file found to force real test or mock a docx
        print("Real file not found. Creating a dummy DOCX...")
        import docx
        doc = docx.Document()
        doc.add_paragraph("This is a test document for summarization. It contains important information about AI and Python.")
        doc.add_paragraph("Key Point 1: AI is powerful.")
        doc.add_paragraph("Key Point 2: Python is versatile.")
        doc.add_paragraph("Conclusion: Combining them is great.")
        doc.save("test_summary.docx")
        filename = "test_summary.docx"
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    with open(filename, "rb") as f:
        file_bytes = f.read()
        
    print("1. Sending to Gemini for Summarization...")
    try:
        summary = summarize_document(file_bytes, mime_type)
        print("Summary Received (Keys):", list(summary.keys()))
        try:
             print(str(summary)[:200] + "...")
        except:
             print("Summary content contains unicode, skipping console print.")
        
        if not summary.get("key_points"):
             print("FAILED: No key points returned.")
             return
             
    except Exception as e:
        # Safe printing for Windows console
        try:
            print(f"Summarization FAILED: {str(e)}")
        except:
            print(f"Summarization FAILED: {str(e).encode('ascii', 'ignore')}")
            
        import traceback
        traceback.print_exc()
        return

    print("\n2. Generating PDF...")
    try:
        output_pdf = "test_summary_output.pdf"
        pdf_path = save_summary_to_pdf(summary, output_pdf)
        
        if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
            print(f"SUCCESS: PDF generated at {pdf_path} (Size: {os.path.getsize(pdf_path)} bytes)")
        else:
            print("FAILED: PDF file empty or not found.")
            
    except Exception as e:
        print(f"PDF Generation FAILED: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    verify_summarizer()
